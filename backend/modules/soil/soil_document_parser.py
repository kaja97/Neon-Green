"""Soil test document parser for multi-format laboratory reports.

Extracts text, numerical data, and structured tables from:
  - PDF documents (pypdf + raw PDF bytes for Gemini Multimodal)
  - Word documents (.docx) (python-docx)
  - Excel spreadsheets (.xlsx, .xls) (openpyxl)
  - Plain CSV files (.csv)
  - Image files (PNG, JPG, WEBP) -> prepared for Gemini Multimodal Vision
"""
import io
import os
import logging
from typing import Tuple, Optional, Any

logger = logging.getLogger(__name__)


def parse_soil_document(
    file_bytes: bytes,
    filename: str,
    content_type: Optional[str] = None,
) -> Tuple[str, Optional[bytes], str]:
    """Extract document text or format multimodal payload for Gemini AI.

    Args:
        file_bytes: Raw bytes of the document
        filename: Name of the uploaded file
        content_type: Optional MIME type from request header

    Returns:
        Tuple of:
          - extracted_text: str (Structured text/markdown from the document)
          - raw_multimodal_bytes: Optional[bytes] (PDF or Image bytes for Gemini Vision)
          - mime_type: str (Resolved MIME type, e.g. "image/png", "application/pdf")
    """
    ext = os.path.splitext(filename.lower())[1]

    # 1. Images (PNG, JPG, JPEG, WEBP)
    if ext in [".png", ".jpg", ".jpeg", ".webp"]:
        resolved_mime = "image/png" if ext == ".png" else "image/webp" if ext == ".webp" else "image/jpeg"
        return "", file_bytes, resolved_mime

    # 2. PDF Documents — send raw PDF bytes to Gemini for high-fidelity multimodal visual analysis
    elif ext == ".pdf":
        extracted_text = _extract_pdf_text(file_bytes)
        return extracted_text, file_bytes, "application/pdf"

    # 3. Word Documents (.docx)
    elif ext in [".docx", ".doc"]:
        text = _extract_docx_text(file_bytes)
        return text, None, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # 4. Excel Spreadsheets (.xlsx, .xls)
    elif ext in [".xlsx", ".xls"]:
        text = _extract_xlsx_text(file_bytes)
        return text, None, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    # 5. CSV Files
    elif ext == ".csv":
        text = _extract_csv_text(file_bytes)
        return text, None, "text/csv"

    else:
        # Fallback text decoding
        try:
            return file_bytes.decode("utf-8", errors="ignore"), None, "text/plain"
        except Exception:
            return "", None, "application/octet-stream"


def _extract_pdf_text(data: bytes) -> str:
    """Extract text and tables from a PDF document using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        text_parts = []

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"--- PAGE {i + 1} ---\n{page_text}")

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning("Failed to extract PDF text with pypdf: %s", e)
        return ""


def _extract_docx_text(data: bytes) -> str:
    """Extract paragraphs and tables from a DOCX document using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        text_parts = []

        # Extract headings and paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for t_idx, table in enumerate(doc.tables):
            table_lines = [f"\n[Table {t_idx + 1}]"]
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            text_parts.append("\n".join(table_lines))

        return "\n".join(text_parts)
    except Exception as e:
        logger.warning("Failed to extract DOCX text: %s", e)
        return ""


def _extract_xlsx_text(data: bytes) -> str:
    """Extract sheet rows and cell values from an Excel spreadsheet using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"\n=== Sheet: {sheet_name} ===")
            for row in sheet.iter_rows(values_only=True):
                # Filter out completely empty rows
                if any(cell is not None and str(cell).strip() != "" for cell in row):
                    line = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    text_parts.append(line)

        return "\n".join(text_parts)
    except Exception as e:
        logger.warning("Failed to extract XLSX text: %s", e)
        return ""


def _extract_csv_text(data: bytes) -> str:
    """Decode and format CSV data."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="ignore")
    return text
