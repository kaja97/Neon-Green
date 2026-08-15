"""Tests for file security scanner, document parsers, and soil AI prompt builders."""
import io
import pytest
from core.file_security import scan_file_safety, FileSecurityError
from modules.soil.soil_document_parser import parse_soil_document
from modules.soil.soil_extractor_prompt import build_soil_extractor_system_prompt
import openpyxl
import docx


# ── File Security Tests ──

class TestFileSecurity:
    def test_empty_file_rejected(self):
        with pytest.raises(FileSecurityError, match="empty"):
            scan_file_safety(b"", "report.pdf")

    def test_unsupported_extension_rejected(self):
        with pytest.raises(FileSecurityError, match="Unsupported file format"):
            scan_file_safety(b"some content", "malware.exe")

    def test_executable_pe_header_rejected(self):
        # Fake MZ executable pretending to be PDF
        with pytest.raises(FileSecurityError, match="Dangerous executable"):
            scan_file_safety(b"MZ\x90\x00\x03fake-exe-content", "trojan.pdf")

    def test_executable_elf_header_rejected(self):
        # Linux ELF header
        with pytest.raises(FileSecurityError, match="Dangerous executable"):
            scan_file_safety(b"\x7fELF\x02\x01\x01\x00fake-elf-binary", "payload.png")

    def test_valid_pdf_passes(self):
        pdf_bytes = b"%PDF-1.4\n1 0 obj\n<< /Title (Soil Test Report) >>\nendobj\n"
        is_safe, msg = scan_file_safety(pdf_bytes, "lab_report.pdf")
        assert is_safe is True
        assert "PDF" in msg

    def test_valid_png_passes(self):
        png_bytes = b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
        is_safe, msg = scan_file_safety(png_bytes, "soil_photo.png")
        assert is_safe is True
        assert "PNG" in msg

    def test_valid_jpg_passes(self):
        jpg_bytes = b"\xff\xd8\xff\xe0\x00\x10JFIF"
        is_safe, msg = scan_file_safety(jpg_bytes, "report.jpg")
        assert is_safe is True
        assert "JPG" in msg

    def test_valid_csv_passes(self):
        csv_bytes = b"Parameter,Value,Unit\npH,6.5,\nNitrogen,240,ppm\n"
        is_safe, msg = scan_file_safety(csv_bytes, "soil_data.csv")
        assert is_safe is True
        assert "CSV" in msg

    def test_null_byte_in_csv_rejected(self):
        with pytest.raises(FileSecurityError, match="null bytes"):
            scan_file_safety(b"pH,6.5\x00malicious", "exploit.csv")


# ── Document Parser Tests ──

class TestDocumentParser:
    def test_parse_csv(self):
        csv_content = b"Parameter,Result,Unit\npH,6.4,\nNitrogen,260,ppm\nPhosphorus,24,ppm\n"
        text, raw_img, mime = parse_soil_document(csv_content, "soil.csv")
        assert "Nitrogen" in text
        assert "6.4" in text
        assert raw_img is None
        assert mime == "text/csv"

    def test_parse_image(self):
        img_bytes = b"\x89PNG\r\n\x1a\n\x00\x00"
        text, raw_img, mime = parse_soil_document(img_bytes, "test.png")
        assert text == ""
        assert raw_img == img_bytes
        assert mime == "image/png"

    def test_parse_xlsx(self):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Soil Results"
        ws.append(["Nutrient", "Concentration", "Unit"])
        ws.append(["pH (1:2.5)", 6.5, ""])
        ws.append(["Available Nitrogen", 280, "ppm"])
        ws.append(["Available Phosphorus", 28, "ppm"])

        buf = io.BytesIO()
        wb.save(buf)
        xlsx_bytes = buf.getvalue()

        text, raw_img, mime = parse_soil_document(xlsx_bytes, "soil_test.xlsx")
        assert "Soil Results" in text
        assert "Available Nitrogen" in text
        assert "280" in text
        assert raw_img is None

    def test_parse_docx(self):
        doc = docx.Document()
        doc.add_heading("Laboratory Soil Analysis", level=1)
        doc.add_paragraph("Sample ID: ST-2026-001")
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "Nutrient"
        table.rows[0].cells[1].text = "Value"
        table.rows[0].cells[2].text = "Unit"
        table.rows[1].cells[0].text = "Potassium (K)"
        table.rows[1].cells[1].text = "190"
        table.rows[1].cells[2].text = "ppm"

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        text, raw_img, mime = parse_soil_document(docx_bytes, "report.docx")
        assert "Laboratory Soil Analysis" in text
        assert "Potassium (K)" in text
        assert "190" in text
        assert raw_img is None


# ── AI Prompt Tests ──

class TestSoilPrompt:
    def test_extractor_prompt_contains_all_schema_fields(self):
        prompt = build_soil_extractor_system_prompt()
        assert "ph_level" in prompt
        assert "electrical_conductivity_ec" in prompt
        assert "organic_carbon_oc" in prompt
        assert "cation_exchange_capacity_cec" in prompt
        assert "nitrogen_n" in prompt
        assert "phosphorus_p" in prompt
        assert "potassium_k" in prompt
        assert "calcium_ca" in prompt
        assert "magnesium_mg" in prompt
        assert "sulfur_s" in prompt
        assert "zinc_zn" in prompt
        assert "boron_b" in prompt
        assert "iron_fe" in prompt
        assert "manganese_mn" in prompt
        assert "copper_cu" in prompt
